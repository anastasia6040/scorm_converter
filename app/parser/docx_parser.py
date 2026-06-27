from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table as DocxTable
# from docx.text.paragraph import Paragraph as DocxParagraph

from .models import (
    ParsedDocument, Heading, Paragraph, ListItem,
    ImageBlock, Table, TableRow, TableCell,
    RunSegment, Alignment, ElementType, Section
)

ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.CENTER: Alignment.CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT: Alignment.RIGHT,
    WD_ALIGN_PARAGRAPH.JUSTIFY: Alignment.JUSTIFY,
    WD_ALIGN_PARAGRAPH.LEFT: Alignment.LEFT,
}

HIGHLIGHT_COLOR_MAP = {
    "YELLOW": "#FFFF99",
    "BRIGHT_GREEN": "#CCFFCC",
    "TURQUOISE": "#CCFFFF",
    "PINK": "#FFB6C1",
    "RED": "#FF9999",
    "DARK_BLUE": "#00008B",
    "TEAL": "#008080",
    "GREEN": "#90EE90",
    "VIOLET": "#EE82EE",
    "DARK_RED": "#8B0000",
    "DARK_YELLOW": "#FFFF66",
    "GRAY_50": "#808080",
    "GRAY_25": "#C0C0C0",
}


ORDERED_NUM_FORMATS = {
    "decimal", "decimalZero", "upperRoman", "lowerRoman",
    "upperLetter", "lowerLetter", "ordinal", "ordinalText", "cardinalText",
}


def _build_numfmt_map(doc) -> dict[str, str]:
    """
    Строит словарь numId -> numFmt (например "11" -> "bullet", "13" -> "decimal"),
    читая word/numbering.xml. numFmt берётся с уровня 0 (ilvl=0), так как
    наши списки в учебных документах не используют многоуровневую нумерацию.

    Возвращает {} если в документе вообще нет numbering.xml (списков нет).
    """
    numfmt_by_numid: dict[str, str] = {}

    numbering_part = None
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return numfmt_by_numid

    if numbering_part is None:
        return numfmt_by_numid

    root = numbering_part.element

    numid_to_abstract: dict[str, str] = {}
    for num_el in root.findall(qn("w:num")):
        num_id = num_el.get(qn("w:numId"))
        abstract_el = num_el.find(qn("w:abstractNumId"))
        if num_id is not None and abstract_el is not None:
            abstract_id = abstract_el.get(qn("w:val"))
            numid_to_abstract[num_id] = abstract_id

    abstract_to_fmt: dict[str, str] = {}
    for abs_el in root.findall(qn("w:abstractNum")):
        abs_id = abs_el.get(qn("w:abstractNumId"))
        if abs_id is None:
            continue
        for lvl_el in abs_el.findall(qn("w:lvl")):
            ilvl = lvl_el.get(qn("w:ilvl"))
            if ilvl == "0":
                fmt_el = lvl_el.find(qn("w:numFmt"))
                if fmt_el is not None:
                    abstract_to_fmt[abs_id] = fmt_el.get(qn("w:val"))
                break

    for num_id, abs_id in numid_to_abstract.items():
        fmt = abstract_to_fmt.get(abs_id)
        if fmt:
            numfmt_by_numid[num_id] = fmt

    return numfmt_by_numid


def _get_num_id(para) -> str | None:
    """Возвращает numId параграфа (ссылку на список нумерации), если он есть."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    numid_el = numPr.find(qn("w:numId"))
    if numid_el is None:
        return None
    return numid_el.get(qn("w:val"))


def _parse_runs(para) -> list[RunSegment]:
    segments = []

    for child in para._p:
        tag = child.tag.split("}")[-1]

        if tag == "r":
            segment = _parse_single_run(child, link=None)
            if segment:
                segments.append(segment)

        elif tag == "hyperlink":
            r_id = child.get(qn("r:id"))
            url = None
            try:
                if r_id and r_id in para.part.rels:
                    url = para.part.rels[r_id].target_ref
            except Exception:
                pass

            for run_el in child.findall(qn("w:r")):
                segment = _parse_single_run(run_el, link=url)
                if segment:
                    segments.append(segment)

    return segments


def _parse_single_run(run_el, link: str | None) -> RunSegment | None:
    """Парсит один XML элемент w:r в RunSegment."""
    t_el = run_el.find(qn("w:t"))
    if t_el is None or not t_el.text:
        return None

    text = t_el.text

    rpr = run_el.find(qn("w:rPr"))

    bold = rpr is not None and rpr.find(qn("w:b")) is not None
    italic = rpr is not None and rpr.find(qn("w:i")) is not None
    underline = rpr is not None and rpr.find(qn("w:u")) is not None

    color = None
    if rpr is not None:
        color_el = rpr.find(qn("w:color"))
        if color_el is not None:
            val = color_el.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
            if val and val.lower() != "auto":
                color = val

    highlight_css = None
    if rpr is not None:
        hl_el = rpr.find(qn("w:highlight"))
        if hl_el is not None:
            name = hl_el.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "").upper()
            highlight_css = HIGHLIGHT_COLOR_MAP.get(name)

    underline = (
        rpr is not None and
        rpr.find(qn("w:u")) is not None and
        rpr.find(qn("w:u")).get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "") not in ("none", "")
    )

    return RunSegment(
        text=text,
        bold=bold,
        italic=italic,
        underline=underline,
        color=color,
        highlight=highlight_css,
        link=link,
    )


def _get_alignment(para) -> Alignment:
    return ALIGNMENT_MAP.get(para.alignment, Alignment.LEFT)


def _get_background(para) -> str | None:
    shd = para._p.find(".//" + qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill and fill.lower() not in ("auto", "ffffff", "none"):
            return fill
    return None


def _has_image(para) -> bool:
    return para._p.find(".//" + qn("a:blip")) is not None


def _is_list_item(para) -> bool:
    if "List" in para.style.name:
        return True
    return _get_num_id(para) is not None


def _is_ordered_list(para, numfmt_by_numid: dict[str, str]) -> bool:
    num_id = _get_num_id(para)
    if num_id is None:
        return False
    fmt = numfmt_by_numid.get(num_id)
    if fmt is None:
        return False
    return fmt in ORDERED_NUM_FORMATS


def _parse_paragraph(para, doc, result, image_counter, used_image_rids, numfmt_by_numid):
    style = para.style.name
    text = para.text.strip()

    if not text and not _has_image(para):
        return image_counter

    if style.startswith("Heading"):
        try:
            level = int(style.split()[-1])
        except ValueError:
            level = 1
        level = min(level, 3)
        if not result.title and level == 1:
            result.title = text
        result.elements.append(Heading(level=level, text=text))

    elif _has_image(para):
        for blip in para._p.findall(".//" + qn("a:blip")):
            r_embed = blip.get(qn("r:embed"))
            if r_embed and r_embed not in used_image_rids:
                used_image_rids.add(r_embed)
                rel = doc.part.rels.get(r_embed)
                if rel and "image" in rel.reltype:
                    image_data = rel.target_part.blob
                    ext = rel.target_part.content_type.split("/")[-1]
                    if ext == "jpeg":
                        ext = "jpg"
                    filename = f"image{image_counter}.{ext}"
                    image_counter += 1
                    result.images[filename] = image_data
                    result.elements.append(ImageBlock(filename=filename))

    elif _is_list_item(para):
        runs = _parse_runs(para)
        if runs:
            result.elements.append(ListItem(
                runs=runs,
                ordered=_is_ordered_list(para, numfmt_by_numid),
            ))

    else:
        runs = _parse_runs(para)
        if runs:
            result.elements.append(Paragraph(
                runs=runs,
                alignment=_get_alignment(para),
                background=_get_background(para),
            ))

    return image_counter


def _parse_table(docx_table) -> Table:
    parsed_table = Table()
    for row in docx_table.rows:
        parsed_row = TableRow()
        for cell in row.cells:
            parsed_row.cells.append(TableCell(text=cell.text.strip()))
        parsed_table.rows.append(parsed_row)
    return parsed_table


def parse_docx(file_path: str) -> ParsedDocument:
    doc = Document(file_path)
    result = ParsedDocument()
    image_counter = 1
    used_image_rids = set()

    numfmt_by_numid = _build_numfmt_map(doc)

    para_map = {p._element: p for p in doc.paragraphs}

    table_map = {t._element: t for t in doc.tables}

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]

        if tag == "p" and child in para_map:
            para = para_map[child]
            image_counter = _parse_paragraph(
                para, doc, result, image_counter, used_image_rids, numfmt_by_numid
            )

        elif tag == "tbl" and child in table_map:
            docx_table = table_map[child]
            result.elements.append(_parse_table(docx_table))

    if not result.title:
        result.title = "Без названия"

    sections = []
    current_section = None

    for element in result.elements:
        if isinstance(element, Heading) and element.level == 1:
            current_section = Section(title=element.text)
            sections.append(current_section)
        elif current_section is None:
            current_section = Section(title=result.title)
            sections.append(current_section)
            current_section.elements.append(element)
        else:
            current_section.elements.append(element)

    result.sections = sections if sections else [
        Section(title=result.title, elements=result.elements)]

    return result
